"""研报导出服务：Markdown → Word (.docx) / PDF。

Word 导出仅依赖 python-docx，无需外部工具。
PDF 导出依赖系统级 wkhtmltopdf；不可用时 graceful 降级并警告。
"""
from __future__ import annotations

import re
from pathlib import Path

import markdown as md
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ..utils.logger import get_logger

logger = get_logger("report-exporter")

_CSS = """
<style>
body { font-family: 'Helvetica Neue', Arial, sans-serif; max-width: 800px; margin: auto; padding: 20px; }
h1 { color: #1a1a2e; border-bottom: 2px solid #16213e; padding-bottom: 8px; }
h2 { color: #16213e; margin-top: 24px; }
table { border-collapse: collapse; width: 100%; margin: 12px 0; }
th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
th { background: #f2f2f2; }
</style>
"""


class ReportExporter:

    def _md_to_html(self, md_text: str) -> str:
        return md.markdown(md_text, extensions=["tables", "fenced_code", "toc"])

    def export_word(self, md_text: str, output_path: str) -> str:
        """Markdown → .docx. Returns output path."""
        doc = Document()
        self._current_table = None  # reset table state

        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        for line in md_text.strip().split("\n"):
            stripped = line.strip()
            if not stripped:
                self._current_table = None  # blank line ends table
                continue

            if stripped.startswith("# "):
                self._current_table = None
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("## "):
                self._current_table = None
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("### "):
                self._current_table = None
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("|") and "---" not in stripped:
                self._add_table_row(doc, stripped)
            elif stripped.startswith(("- ", "* ", "1. ")):
                self._current_table = None
                text = re.sub(r"^[\-\*]\s+|\d+\.\s+", "", stripped)
                text = self._strip_inline_md(text)
                doc.add_paragraph(text, style="List Bullet")
            else:
                self._current_table = None
                text = self._strip_inline_md(stripped)
                if text:
                    doc.add_paragraph(text)

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        logger.info(f"Word report saved: {output_path}")
        return output_path

    def export_pdf(self, md_text: str, output_path: str) -> str | None:
        """Markdown → PDF. Returns None if wkhtmltopdf unavailable."""
        try:
            import pdfkit
        except ImportError:
            logger.warning("pdfkit not installed, PDF export unavailable")
            return None

        html = f"<html><head>{_CSS}</head><body>{self._md_to_html(md_text)}</body></html>"
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)

        try:
            pdfkit.from_string(html, output_path, options={"encoding": "UTF-8", "quiet": ""})
            logger.info(f"PDF report saved: {output_path}")
            return output_path
        except Exception as exc:
            logger.warning(f"PDF export failed (wkhtmltopdf installed?): {exc}")
            return None

    def _add_table_row(self, doc: Document, line: str) -> None:
        cells = [c.strip() for c in line.strip("|").split("|")]
        if self._current_table is None:
            self._current_table = doc.add_table(rows=0, cols=len(cells))
            self._current_table.style = "Table Grid"
        row = self._current_table.add_row()
        for i, cell in enumerate(cells):
            if i < len(row.cells):
                row.cells[i].text = self._strip_inline_md(cell)

    @staticmethod
    def _strip_inline_md(text: str) -> str:
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
        text = re.sub(r"\*(.+?)\*", r"\1", text)
        text = re.sub(r"`(.+?)`", r"\1", text)
        text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
        return text.strip()
